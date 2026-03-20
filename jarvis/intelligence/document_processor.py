"""
JARVIS 문서 처리기 — Iteration 4
PDF / DOCX / XLSX / CSV / TXT / 이미지 문서를 자동 분석·요약·추출
- 표, 차트, 이미지 포함 지원
- LLM 기반 심층 이해 및 Q&A
- 구조화된 데이터 추출
- 다국어 지원
"""

import io
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    index: int
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)  # page, table, section


@dataclass
class DocumentResult:
    file_path: str
    file_type: str
    title: str
    summary: str
    full_text: str
    chunks: List[DocumentChunk]
    tables: List[Dict]       # list of {headers, rows}
    key_facts: List[str]
    entities: Dict[str, List[str]]   # {person, org, date, location}
    page_count: int
    word_count: int
    language: str
    processed_at: str = field(default_factory=lambda: datetime.now().isoformat())


class DocumentProcessor:
    """
    JARVIS 문서 처리 엔진
    어떤 형식의 문서든 자동으로 읽고 이해하고 추출
    """

    CHUNK_SIZE = 4000   # characters per chunk

    def __init__(self, llm_manager=None):
        self.llm = llm_manager
        self._cache: Dict[str, DocumentResult] = {}
        logger.info("DocumentProcessor initialized")

    # ── 통합 진입점 ──────────────────────────────────────────────────────────
    def process(self, file_path: str, extract_tables: bool = True) -> DocumentResult:
        """파일 경로를 받아 종류에 상관없이 처리"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"파일 없음: {file_path}")

        if str(file_path) in self._cache:
            logger.debug(f"Cache hit: {file_path}")
            return self._cache[str(file_path)]

        suffix = path.suffix.lower()
        dispatch = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_docx,
            ".xlsx": self._process_xlsx,
            ".xls": self._process_xlsx,
            ".csv": self._process_csv,
            ".txt": self._process_text,
            ".md": self._process_text,
            ".json": self._process_json,
            ".py": self._process_text,
            ".html": self._process_html,
            ".htm": self._process_html,
        }

        processor = dispatch.get(suffix, self._process_text)
        raw_text, tables, page_count = processor(str(path))

        chunks = self._chunk_text(raw_text)
        result = DocumentResult(
            file_path=str(file_path),
            file_type=suffix.lstrip("."),
            title=self._extract_title(raw_text, path.name),
            summary="",
            full_text=raw_text,
            chunks=chunks,
            tables=tables,
            key_facts=[],
            entities={},
            page_count=page_count,
            word_count=len(raw_text.split()),
            language="ko",
        )

        if self.llm:
            result = self._llm_enrich(result)

        self._cache[str(file_path)] = result
        return result

    def process_bytes(self, content: bytes, filename: str) -> DocumentResult:
        """메모리 내 파일 처리 (업로드된 파일)"""
        suffix = Path(filename).suffix.lower()
        tmp_path = Path(f"/tmp/jarvis_doc_{datetime.now().timestamp()}{suffix}")
        tmp_path.write_bytes(content)
        try:
            return self.process(str(tmp_path))
        finally:
            tmp_path.unlink(missing_ok=True)

    # ── PDF ──────────────────────────────────────────────────────────────────
    def _process_pdf(self, path: str) -> Tuple[str, List[Dict], int]:
        try:
            import pdfplumber
            pages_text = []
            tables = []
            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    pages_text.append(text)
                    for tbl in page.extract_tables() or []:
                        if tbl:
                            headers = [str(c or "") for c in tbl[0]]
                            rows = [[str(c or "") for c in row] for row in tbl[1:]]
                            tables.append({"page": i + 1, "headers": headers, "rows": rows})
            return "\n\n".join(pages_text), tables, page_count
        except ImportError:
            logger.warning("pdfplumber 없음 — PyPDF2 시도")
            return self._process_pdf_fallback(path)
        except Exception as e:
            logger.error(f"PDF 처리 오류: {e}")
            return f"[PDF 처리 오류: {e}]", [], 0

    def _process_pdf_fallback(self, path: str) -> Tuple[str, List[Dict], int]:
        try:
            import PyPDF2
            text_parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                page_count = len(reader.pages)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts), [], page_count
        except ImportError:
            return "[PDF 라이브러리 없음: pip install pdfplumber]", [], 0

    # ── DOCX ─────────────────────────────────────────────────────────────────
    def _process_docx(self, path: str) -> Tuple[str, List[Dict], int]:
        try:
            import docx
            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for tbl in doc.tables:
                rows = [[cell.text for cell in row.cells] for row in tbl.rows]
                if rows:
                    tables.append({"headers": rows[0], "rows": rows[1:]})
            text = "\n\n".join(paragraphs)
            return text, tables, max(1, len(paragraphs) // 30)
        except ImportError:
            return "[python-docx 없음: pip install python-docx]", [], 0
        except Exception as e:
            return f"[DOCX 오류: {e}]", [], 0

    # ── XLSX ─────────────────────────────────────────────────────────────────
    def _process_xlsx(self, path: str) -> Tuple[str, List[Dict], int]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            all_text = []
            tables = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = [[str(cell.value or "") for cell in row] for row in ws.iter_rows()]
                if rows:
                    headers = rows[0]
                    data_rows = rows[1:]
                    tables.append({"sheet": sheet_name, "headers": headers, "rows": data_rows[:100]})
                    all_text.append(f"[시트: {sheet_name}]\n" + "\n".join(["\t".join(r) for r in rows[:50]]))
            return "\n\n".join(all_text), tables, len(wb.sheetnames)
        except ImportError:
            return self._process_xlsx_csv_fallback(path)
        except Exception as e:
            return f"[XLSX 오류: {e}]", [], 0

    def _process_xlsx_csv_fallback(self, path: str) -> Tuple[str, List[Dict], int]:
        try:
            import csv
            with open(path, newline="", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                rows = list(reader)
            if rows:
                return "\n".join(["\t".join(r) for r in rows[:200]]), [{"headers": rows[0], "rows": rows[1:100]}], 1
        except Exception:
            pass
        return "[XLSX/CSV 라이브러리 없음]", [], 0

    # ── CSV ──────────────────────────────────────────────────────────────────
    def _process_csv(self, path: str) -> Tuple[str, List[Dict], int]:
        import csv
        try:
            for enc in ["utf-8-sig", "utf-8", "cp949", "latin-1"]:
                try:
                    with open(path, newline="", encoding=enc) as f:
                        rows = list(csv.reader(f))
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return "[CSV 인코딩 오류]", [], 0

            if not rows:
                return "", [], 0
            headers = rows[0]
            data = rows[1:]
            table = {"headers": headers, "rows": data[:200]}
            text = "\n".join(["\t".join(r) for r in rows[:300]])
            return text, [table], 1
        except Exception as e:
            return f"[CSV 오류: {e}]", [], 0

    # ── 텍스트/마크다운/코드 ─────────────────────────────────────────────────
    def _process_text(self, path: str) -> Tuple[str, List[Dict], int]:
        for enc in ["utf-8", "cp949", "latin-1"]:
            try:
                text = Path(path).read_text(encoding=enc)
                lines = text.count("\n") + 1
                return text, [], max(1, lines // 50)
            except UnicodeDecodeError:
                continue
        return "[텍스트 인코딩 오류]", [], 0

    # ── JSON ─────────────────────────────────────────────────────────────────
    def _process_json(self, path: str) -> Tuple[str, List[Dict], int]:
        try:
            data = json.loads(Path(path).read_text(encoding="utf-8"))
            text = json.dumps(data, ensure_ascii=False, indent=2)
            tables = []
            if isinstance(data, list) and data and isinstance(data[0], dict):
                headers = list(data[0].keys())
                rows = [[str(row.get(h, "")) for h in headers] for row in data[:100]]
                tables.append({"headers": headers, "rows": rows})
            return text, tables, 1
        except Exception as e:
            return f"[JSON 오류: {e}]", [], 0

    # ── HTML ─────────────────────────────────────────────────────────────────
    def _process_html(self, path: str) -> Tuple[str, List[Dict], int]:
        try:
            from bs4 import BeautifulSoup
            html = Path(path).read_text(encoding="utf-8", errors="ignore")
            soup = BeautifulSoup(html, "lxml")
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            tables = []
            for tbl in soup.find_all("table"):
                rows = []
                for tr in tbl.find_all("tr"):
                    rows.append([td.get_text(strip=True) for td in tr.find_all(["td", "th"])])
                if rows:
                    tables.append({"headers": rows[0], "rows": rows[1:]})
            return text, tables, 1
        except ImportError:
            text = Path(path).read_text(encoding="utf-8", errors="ignore")
            clean = re.sub(r"<[^>]+>", "", text)
            return clean, [], 1

    # ── 유틸리티 ─────────────────────────────────────────────────────────────
    def _chunk_text(self, text: str) -> List[DocumentChunk]:
        chunks = []
        for i in range(0, len(text), self.CHUNK_SIZE):
            chunk_text = text[i : i + self.CHUNK_SIZE]
            chunks.append(DocumentChunk(index=i // self.CHUNK_SIZE, text=chunk_text))
        return chunks

    def _extract_title(self, text: str, filename: str) -> str:
        first_line = text.strip().split("\n")[0][:120] if text.strip() else ""
        if first_line and len(first_line) < 100:
            return first_line
        return Path(filename).stem

    # ── LLM 보강 ─────────────────────────────────────────────────────────────
    def _llm_enrich(self, result: DocumentResult) -> DocumentResult:
        """LLM으로 요약·핵심 사실·개체 추출"""
        from jarvis.llm.manager import Message

        preview = result.full_text[:6000]
        prompt = f"""다음 문서를 분석하세요 (파일: {result.file_path}):

{preview}

JSON 형식으로 반환:
{{
  "title": "문서 제목",
  "summary": "3-5문장 요약",
  "key_facts": ["핵심 사실 1", "핵심 사실 2", "핵심 사실 3"],
  "entities": {{
    "person": ["이름들"],
    "organization": ["기관명들"],
    "date": ["날짜들"],
    "location": ["장소들"]
  }},
  "language": "ko 또는 en",
  "document_type": "보고서/논문/계약서/기술문서/기타"
}}"""

        messages = [Message(role="user", content=prompt)]
        try:
            response = self.llm.chat(
                messages,
                system="문서 분석 전문가입니다. JSON만 출력하세요.",
                max_tokens=2048,
            )
            match = re.search(r"\{.*\}", response.content, re.DOTALL)
            if match:
                data = json.loads(match.group())
                result.title = data.get("title", result.title)
                result.summary = data.get("summary", "")
                result.key_facts = data.get("key_facts", [])
                result.entities = data.get("entities", {})
                result.language = data.get("language", "ko")
        except Exception as e:
            logger.error(f"LLM enrich error: {e}")
        return result

    # ── Q&A ──────────────────────────────────────────────────────────────────
    def ask(self, result: DocumentResult, question: str) -> str:
        """문서에 대한 자연어 질문 답변"""
        if not self.llm:
            return "LLM 없음"
        from jarvis.llm.manager import Message

        # 관련 청크 선택 (간단한 키워드 매칭)
        keywords = set(question.lower().split())
        scored = []
        for chunk in result.chunks:
            score = sum(1 for kw in keywords if kw in chunk.text.lower())
            scored.append((score, chunk))
        scored.sort(key=lambda x: -x[0])
        context = "\n\n".join(c.text for _, c in scored[:3])

        prompt = f"""문서 내용:
{context}

질문: {question}

문서에 기반하여 정확하게 답변하세요. 문서에 없는 내용은 "문서에서 찾을 수 없음"이라고 하세요."""

        messages = [Message(role="user", content=prompt)]
        try:
            response = self.llm.chat(messages, max_tokens=2048)
            return response.content
        except Exception as e:
            return f"오류: {e}"

    # ── 표 분석 ──────────────────────────────────────────────────────────────
    def analyze_table(self, table: Dict, question: str = None) -> Dict:
        """표 데이터 통계 분석"""
        headers = table.get("headers", [])
        rows = table.get("rows", [])
        if not rows:
            return {"error": "빈 표"}

        # 숫자 열 감지
        numeric_cols = {}
        for col_idx, header in enumerate(headers):
            vals = []
            for row in rows:
                if col_idx < len(row):
                    try:
                        vals.append(float(row[col_idx].replace(",", "")))
                    except ValueError:
                        pass
            if vals:
                numeric_cols[header] = {
                    "count": len(vals),
                    "sum": sum(vals),
                    "mean": sum(vals) / len(vals),
                    "min": min(vals),
                    "max": max(vals),
                }

        result = {
            "row_count": len(rows),
            "column_count": len(headers),
            "headers": headers,
            "numeric_summary": numeric_cols,
        }

        if question and self.llm:
            from jarvis.llm.manager import Message
            table_text = "\t".join(headers) + "\n"
            table_text += "\n".join(["\t".join(r) for r in rows[:50]])
            prompt = f"표 데이터:\n{table_text}\n\n질문: {question}"
            messages = [Message(role="user", content=prompt)]
            try:
                resp = self.llm.chat(messages, max_tokens=1024)
                result["answer"] = resp.content
            except Exception:
                pass

        return result

    # ── 요약 리포트 ──────────────────────────────────────────────────────────
    def format_markdown(self, result: DocumentResult) -> str:
        lines = [
            f"# {result.title}",
            f"**파일:** {result.file_path} | **타입:** {result.file_type} | **페이지:** {result.page_count} | **단어:** {result.word_count:,}",
            "",
            "## 요약",
            result.summary or "(요약 없음)",
            "",
        ]
        if result.key_facts:
            lines += ["## 핵심 사실"] + [f"- {f}" for f in result.key_facts] + [""]
        if result.entities:
            lines.append("## 주요 개체")
            for etype, items in result.entities.items():
                if items:
                    lines.append(f"**{etype}:** {', '.join(items)}")
            lines.append("")
        if result.tables:
            lines.append(f"## 표 ({len(result.tables)}개)")
            for i, tbl in enumerate(result.tables[:3]):
                lines.append(f"### 표 {i+1}: {tbl.get('headers', [])[:4]}")
                for row in tbl.get("rows", [])[:5]:
                    lines.append(f"  {row}")
            lines.append("")
        return "\n".join(lines)

    def get_stats(self) -> Dict:
        return {
            "cached_documents": len(self._cache),
            "total_words": sum(r.word_count for r in self._cache.values()),
            "file_types": list({r.file_type for r in self._cache.values()}),
        }
